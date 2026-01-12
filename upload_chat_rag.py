from flask import Flask, render_template_string, jsonify, request, send_file
import sys
import json
import os
import threading
from pathlib import Path
from dotenv import load_dotenv
from datetime import datetime
from openai import OpenAI
import uuid

from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import torch
import re

# Document processing imports
from pypdf import PdfReader
try:
    from docx import Document as DocxDocument
except ImportError:
    # Fallback if python-docx is not installed
    DocxDocument = None
import io

# Force unbuffered output
sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

load_dotenv(override=True)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = 'uploaded_documents'

# Create upload folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# ============================================================================
# CONFIGURATION
# ============================================================================

# LLM Configuration - Using Ollama
env_path = r"C:\Users\Cem\Desktop\projeler\YapayZekaUygulamalari\.env"
load_dotenv(env_path)

# Ollama Configuration
ollama_base_url = "http://localhost:11434/v1"
ollama_model = "llama3.2"
ollama_llm = None

try:
    ollama_llm = ChatOllama(temperature=0, model=ollama_model)
    print("✅ Ollama client ready")
except Exception as e:
    print(f"⚠️  Ollama not available: {e}")

# RAG Configuration
EMBEDDING_MODEL = "intfloat/multilingual-e5-large"
device = 'cpu'
RETRIEVAL_K = 10
db_name = "vector_db_uploaded_faiss"

HOLY_PROMPT = """You are a knowledgeable, friendly assistant.
You are chatting with a user about documents they have uploaded.
If relevant, use the given context from the uploaded documents to answer any question.
If you don't know the answer, say so.
Context:
{context}
"""
print("=" * 60)
print("💬 Upload RAG Chatbot - Initializing")
print("=" * 60)

# Initialize embeddings
print(f"\n⏳ Loading embedding model: {EMBEDDING_MODEL}...")
embeddings = HuggingFaceEmbeddings(
    model_name=EMBEDDING_MODEL,
    model_kwargs={'device': device},
    encode_kwargs={'normalize_embeddings': True}
)
print("✅ Embedding model loaded")

# Load or create vectorstore
print(f"\n💾 Loading vectorstore: {db_name}...")
vector_store = None
if os.path.exists(db_name) and os.path.exists(os.path.join(db_name, "index.faiss")):
    try:
        vector_store = FAISS.load_local(
            folder_path=db_name,
            embeddings=embeddings,
            allow_dangerous_deserialization=True
        )
        doc_count = len(vector_store.docstore._dict)
        print(f"✅ Vectorstore loaded ({doc_count} documents)")
    except Exception as e:
        print(f"⚠️  Error loading vectorstore: {e}. Creating new one.")
        vector_store = FAISS.from_texts([""], embeddings)
        vector_store.save_local(db_name)
        print("✅ New vectorstore created")
else:
    print("📝 Creating new vectorstore...")
    vector_store = FAISS.from_texts([""], embeddings)
    vector_store.save_local(db_name)
    print("✅ New vectorstore created")

# Thread lock for vectorstore operations
_vectorstore_lock = threading.Lock()

# Text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    length_function=len,
)

print("\n✅ Upload RAG System ready!")
print("=" * 60)

def process_pdf(file_content: bytes, filename: str) -> list[Document]:
    """Extract text from PDF file"""
    documents = []
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        full_text = ""
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                full_text += f"\n\n--- Page {page_num + 1} ---\n\n{text}"
        
        if full_text.strip():
            # Split into chunks
            chunks = text_splitter.split_text(full_text)
            
            for idx, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "file_type": "pdf",
                        "chunk_index": idx,
                        "total_chunks": len(chunks),
                        "upload_date": datetime.now().isoformat()
                    }
                )
                documents.append(doc)
        
        print(f"✅ Processed PDF: {filename} - {len(documents)} chunks")
    except Exception as e:
        print(f"❌ Error processing PDF {filename}: {e}")
        raise
    
    return documents

def process_docx(file_content: bytes, filename: str) -> list[Document]:
    """Extract text from Word document"""
    if DocxDocument is None:
        raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")
    
    documents = []
    try:
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        
        full_text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                full_text += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    full_text += row_text + "\n"
        
        if full_text.strip():
            # Split into chunks
            chunks = text_splitter.split_text(full_text)
            
            for idx, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "file_type": "docx",
                        "chunk_index": idx,
                        "total_chunks": len(chunks),
                        "upload_date": datetime.now().isoformat()
                    }
                )
                documents.append(doc)
        
        print(f"✅ Processed DOCX: {filename} - {len(documents)} chunks")
    except Exception as e:
        print(f"❌ Error processing DOCX {filename}: {e}")
        raise
    
    return documents

def process_txt(file_content: bytes, filename: str) -> list[Document]:
    """Extract text from plain text file"""
    documents = []
    try:
        text = file_content.decode('utf-8', errors='ignore')
        
        if text.strip():
            # Split into chunks
            chunks = text_splitter.split_text(text)
            
            for idx, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "file_type": "txt",
                        "chunk_index": idx,
                        "total_chunks": len(chunks),
                        "upload_date": datetime.now().isoformat()
                    }
                )
                documents.append(doc)
        
        print(f"✅ Processed TXT: {filename} - {len(documents)} chunks")
    except Exception as e:
        print(f"❌ Error processing TXT {filename}: {e}")
        raise
    
    return documents

def process_uploaded_file(file) -> list[Document]:
    """Process uploaded file based on its extension"""
    filename = file.filename
    file_content = file.read()
    
    # Determine file type
    file_ext = Path(filename).suffix.lower()
    
    if file_ext == '.pdf':
        return process_pdf(file_content, filename)
    elif file_ext in ['.docx', '.doc']:
        return process_docx(file_content, filename)
    elif file_ext == '.txt':
        return process_txt(file_content, filename)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

def add_documents_to_vectorstore(documents: list[Document]):
    """Add documents to the vector store"""
    if not documents:
        return
    
    try:
        with _vectorstore_lock:
            # Add documents to existing vectorstore
            vector_store.add_documents(documents)
            # Save the updated vectorstore
            vector_store.save_local(db_name)
        
        print(f"✅ Added {len(documents)} document chunks to vectorstore")
    except Exception as e:
        print(f"❌ Error adding documents to vectorstore: {e}")
        raise
   

def expand_query(question: str) -> list[str]:
    """Expand query with synonyms and related terms for better retrieval."""
    question_lower = question.lower()
    expanded_queries = [question]  # Always include original
    
    # Company/employer related
    if any(word in question_lower for word in ['company', 'employer', 'work', 'working', 'job', 'current']):
        if 'company' in question_lower:
            expanded_queries.append(question.replace('company', 'employer'))
            expanded_queries.append(question.replace('company', 'workplace'))
        if 'current' in question_lower:
            expanded_queries.append(question.replace('current', 'present'))
            expanded_queries.append(question.replace('currently', 'now'))
    
    # Extract key terms
    key_terms = []
    important_words = ['company', 'employer', 'work', 'working', 'current', 'name', 'position', 'role', 'job']
    for word in important_words:
        if word in question_lower:
            key_terms.append(word)
    
    # Create focused queries with key terms
    if key_terms:
        focused_query = " ".join(key_terms)
        if focused_query not in expanded_queries:
            expanded_queries.append(focused_query)
    
    return expanded_queries

def fetch_context(question: str, use_mmr: bool = True) -> list[Document]:
    """Retrieve relevant context documents for a question with improved retrieval."""
    try:
        with _vectorstore_lock:
            all_docs = []
            seen_content = set()
            
            # Try multiple query variations
            expanded_queries = expand_query(question)
            
            for query in expanded_queries[:3]:  # Limit to top 3 variations
                try:
                    if use_mmr:
                        docs = vector_store.max_marginal_relevance_search(
                            query, 
                            k=RETRIEVAL_K,
                            fetch_k=min(RETRIEVAL_K * 3, 30)  # Increased fetch_k for better diversity
                        )
                    else:
                        docs = vector_store.similarity_search(query, k=RETRIEVAL_K)
                    
                    # Deduplicate by content
                    for doc in docs:
                        content_hash = hash(doc.page_content[:100])  # Hash first 100 chars
                        if content_hash not in seen_content:
                            seen_content.add(content_hash)
                            all_docs.append(doc)
                except Exception as e:
                    print(f"⚠️  Error with query '{query}': {e}", flush=True)
                    continue
            
            # Sort by relevance (if we have scores, otherwise keep order)
            # Limit to top K documents
            final_docs = all_docs[:RETRIEVAL_K * 2]  # Get more candidates
            
            print(f"🔍 Retrieved {len(final_docs)} unique documents from {len(expanded_queries)} query variations", flush=True)
            return final_docs
    except Exception as e:
        print(f"❌ Error in fetch_context: {str(e)}", flush=True)
        return []

def rewrite_query_with_context(question: str, history: list[dict] = []) -> str:
    """Rewrite the current question to be self-contained using minimal context from history."""
    if not history:
        return question
    
    recent_context = []
    for msg in history[-4:]:
        if msg.get("role") == "user":
            recent_context.append(f"Previous question: {msg.get('content', '')}")
        elif msg.get("role") == "assistant":
            recent_context.append(f"Previous answer: {msg.get('content', '')[:200]}")
    
    pronouns = ["it", "its", "this", "that", "they", "them", "these", "those", "the", "which"]
    question_lower = question.lower()
    needs_context = any(pronoun in question_lower for pronoun in pronouns) or len(question.split()) < 5
    
    if needs_context and recent_context:
        context_str = " ".join(recent_context[-2:])
        rewritten = f"{context_str} {question}"
        if len(rewritten) > 200:
            rewritten = question
        return rewritten
    else:
        return question

search_knowledge_base_json = {
    "name": "search_knowledge_base",
    "description": "Search the uploaded documents knowledge base for information. Use this when you need factual information from the uploaded documents.",
    "parameters": {
        "type": "object",
        "properties": {
            "query": {
                "type": "string",
                "description": "The search query to find relevant information in the knowledge base"
            }
        },
        "required": ["query"],
        "additionalProperties": False
    }
}

tools = [
    {"type": "function", "function": search_knowledge_base_json}
]

def create_ollama_tool_prompt(question: str, history: list[dict] = [], tool_results: str = None):
    """Create a prompt that instructs Ollama to use tools."""
    
    available_tools = """
AVAILABLE TOOLS:

1. search_knowledge_base(query: str)
   - Use this FIRST for questions about information in uploaded documents
   - IMPORTANT: When searching, use specific keywords from the question
   - For "current company" questions, search for: "company", "employer", "work", "current job"
   - Example: {{"name": "search_knowledge_base", "arguments": {{"query": "current company employer workplace"}}}}

TOOL CALLING FORMAT:
If you need a tool, output ONLY this JSON (no other text):
{{"name": "tool_name", "arguments": {{"arg": "value"}}}}
"""
    
    # Build context from history
    context_str = ""
    if history:
        context_str = "\nPrevious conversation:\n"
        for msg in history[-4:]:
            role = msg.get("role", "")
            content = msg.get("content", "")
            if role == "user":
                context_str += f"User: {content}\n"
            elif role == "assistant":
                context_str += f"Assistant: {content}\n"
    
    # Format tool results with better emphasis on finding exact answers
    if tool_results:
        try:
            result_data = json.loads(tool_results)
            if isinstance(result_data, dict) and "results" in result_data:
                # Format results more clearly
                results_text = ""
                for idx, res in enumerate(result_data.get("results", [])[:5], 1):  # Limit to top 5
                    content = res.get("content", "")[:500]  # Limit content length
                    source = res.get("source", "unknown")
                    results_text += f"\n--- Result {idx} (from {source}) ---\n{content}\n"
                
                formatted_result = f"""
TOOL RESULT RECEIVED - READ CAREFULLY:

Original Question: {result_data.get('original_question', question)}

{results_text}

CRITICAL INSTRUCTIONS:
1. Read ALL the results above carefully
2. Find the EXACT answer to the question: "{question}"
3. If the question asks for a "company name", look for company names, employer names, or workplace names
4. If the question asks "what is the name of the company", extract ONLY the company name, not other names
5. Answer directly and concisely - do NOT include explanations unless asked
6. DO NOT call any more tools - you have all the information you need
"""
            else:
                formatted_result = f"""
TOOL RESULT RECEIVED:
{json.dumps(result_data, indent=2, ensure_ascii=False)[:2000]}

CRITICAL: Read the results above and find the EXACT answer to: "{question}"
Answer directly - do NOT call more tools.
"""
        except:
            formatted_result = f"""
TOOL RESULT RECEIVED:
{tool_results[:2000]}

CRITICAL: Based on the results above, answer the question: "{question}"
Answer directly and concisely.
"""
        
        tool_results_str = formatted_result
    else:
        # Improve initial query generation
        question_lower = question.lower()
        query_hints = ""
        
        if any(word in question_lower for word in ['company', 'employer', 'work']):
            query_hints = "\nQUERY HINT: For company/employer questions, search for terms like: 'company', 'employer', 'workplace', 'current job', 'working at'"
        elif any(word in question_lower for word in ['name']):
            query_hints = "\nQUERY HINT: For name questions, search for the specific type of name mentioned (company name, person name, etc.)"
        
        tool_results_str = f"""
PRIORITY: This is a knowledge question. Use search_knowledge_base FIRST to find information in the uploaded documents.
{query_hints}
IMPORTANT: When calling search_knowledge_base, extract key terms from the question and use them in your search query.
"""
    
    prompt = f"""You are an AI assistant helping users with questions about documents they have uploaded.

{available_tools}

{context_str}

User's question: {question}
{tool_results_str}

ANSWERING INSTRUCTIONS:
- Read the retrieved documents VERY carefully
- Find the EXACT information that answers the question
- If asked for a "company name", extract ONLY the company/employer name, not person names
- If asked "what is the name of the company I am currently working in", look for sections about current employment, work experience, or job information
- Answer directly and concisely - just provide the requested information
- If the information is not found, say "I couldn't find this information in the uploaded documents"

DECISION:
- If you need more information → Output JSON tool call
- If you have enough information → Answer the question directly (NO JSON, just your answer)"""
    
    return prompt

def parse_ollama_tool_call(response_text: str):
    """Parse tool call from Ollama response (manual tool calling)."""
    response_text = response_text.strip()
    
    # Check if this looks like JSON (tool call)
    if response_text.startswith("{") or response_text.startswith("{{"):
        pass
    else:
        # Check if this looks like natural language
        natural_language_patterns = [
            r'^The\s+\w+\s+',
            r'^For\s+',
            r'^Based\s+on',
            r'^According\s+to',
            r'^I\s+',
        ]
        
        is_natural_language = any(re.match(pattern, response_text, re.IGNORECASE) for pattern in natural_language_patterns)
        
        if is_natural_language and len(response_text) > 50:
            json_patterns = [
                r'```json\s*(\{.*?\})\s*```',
                r'```\s*(\{.*?"name".*?\})\s*```',
            ]
            for pattern in json_patterns:
                match = re.search(pattern, response_text, re.DOTALL | re.IGNORECASE)
                if match:
                    try:
                        tool_call = json.loads(match.group(1))
                        if "name" in tool_call or "tool" in tool_call:
                            return tool_call
                    except:
                        continue
            return None
    
    # Look for explicit JSON tool call patterns
    patterns = [
        r'```json\s*(\{.*?\})\s*```',
        r'```json\s*(\{\{.*?\}\})\s*```',
        r'```\s*(\{.*?"name".*?"arguments".*?\})\s*```',
        r'^\s*\{[^{]*"name"[^{]*"arguments"[^{]*\}\s*$',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, response_text, re.DOTALL | re.IGNORECASE)
        for match in matches:
            try:
                json_str = match.group(1) if match.groups() else match.group(0)
                json_str = json_str.replace("{{", "{").replace("}}", "}")
                tool_call = json.loads(json_str)
                
                if "name" in tool_call:
                    return tool_call
                elif "tool" in tool_call:
                    return {"name": tool_call["tool"], "arguments": tool_call.get("arguments", {})}
            except json.JSONDecodeError:
                continue
    
    # Try parsing entire response as JSON
    looks_like_json = (response_text.startswith("{") or response_text.startswith("{{")) and len(response_text) < 500
    
    if looks_like_json:
        try:
            json_str = response_text.replace("{{", "{").replace("}}", "}")
            tool_call = json.loads(json_str)
            if "name" in tool_call or "tool" in tool_call:
                if "name" in tool_call:
                    return tool_call
                else:
                    return {"name": tool_call["tool"], "arguments": tool_call.get("arguments", {})}
        except:
            pass
    
    return None

def answer_question_with_ollama(question: str, history: list[dict] = []) -> tuple[str, list[Document]]:
    """Answer question using Ollama with manual tool calling."""
    if not ollama_llm:
        return "Ollama is not available. Please make sure Ollama is running.", []
    
    docs = []
    max_iterations = 10
    iteration = 0
    
    original_question = question
    tool_results = None
    tools_used = []
    
    # Auto-detect if this is a knowledge question that needs RAG
    knowledge_keywords = ['what', 'who', 'where', 'when', 'which', 'name', 'company', 'employer', 'work', 'current']
    is_knowledge_question = any(keyword in question.lower() for keyword in knowledge_keywords)
    
    # For knowledge questions, automatically do an initial search
    if is_knowledge_question and iteration == 0:
        print(f"🔍 Auto-detected knowledge question, performing initial search...", flush=True)
        try:
            # Create optimized search query
            question_lower = question.lower()
            search_terms = []
            
            # Extract important terms
            if 'company' in question_lower or 'employer' in question_lower:
                search_terms.extend(['company', 'employer', 'workplace', 'working at'])
            if 'current' in question_lower:
                search_terms.extend(['current', 'present', 'now'])
            if 'name' in question_lower:
                search_terms.append('name')
            
            # Use expanded query or original
            search_query = " ".join(search_terms) if search_terms else question
            
            docs = fetch_context(search_query, use_mmr=True)
            if docs:
                result = {
                    "query": search_query,
                    "original_question": question,
                    "results": [
                        {
                            "content": doc.page_content,
                            "source": doc.metadata.get("source", "unknown"),
                            "metadata": doc.metadata
                        } for doc in docs
                    ],
                    "count": len(docs),
                    "message": f"Found {len(docs)} relevant document chunks. Read them carefully to find the exact answer."
                }
                tool_results = json.dumps(result, default=str, ensure_ascii=False)
                tools_used.append("search_knowledge_base")
                iteration = 1  # Skip first iteration since we already have results
                print(f"✅ Auto-retrieved {len(docs)} documents", flush=True)
        except Exception as e:
            print(f"⚠️  Auto-search failed: {e}, continuing with normal flow", flush=True)
    
    while iteration < max_iterations:
        iteration += 1
        
        # Create prompt with tool instructions
        prompt = create_ollama_tool_prompt(original_question, history, tool_results)
        messages = [HumanMessage(content=prompt)]
        
        try:
            response = ollama_llm.invoke(messages)
            response_text = response.content.strip()
            
            print(f"🤖 Ollama response (iteration {iteration}): {response_text[:200]}...", flush=True)
            
            # Try to parse tool call
            tool_call = parse_ollama_tool_call(response_text)
            
            if tool_call:
                tool_name = tool_call.get("name", "")
                tool_args = tool_call.get("arguments", {})
                
                if isinstance(tool_args, str):
                    try:
                        tool_args = json.loads(tool_args)
                    except:
                        tool_args = {}
                
                # Check for infinite loops
                if len(tools_used) >= 2 and tools_used[-1] == tool_name and tools_used[-2] == tool_name:
                    print(f"⚠️  Loop detected: {tool_name} called repeatedly. Forcing answer.", flush=True)
                    if tool_results:
                        try:
                            result_data = json.loads(tool_results)
                            answer_prompt = f"""Based on this data, answer the user's question: {original_question}

Data: {json.dumps(result_data, indent=2, ensure_ascii=False)[:1500]}

Provide a clear, direct answer. Do not use JSON format. Just answer the question naturally."""
                            answer_msg = [HumanMessage(content=answer_prompt)]
                            final_response = ollama_llm.invoke(answer_msg)
                            return final_response.content, docs
                        except Exception as e:
                            print(f"⚠️  Error forcing answer: {e}", flush=True)
                    return "Based on the data retrieved, I found information but reached the iteration limit. Please try rephrasing your question.", docs
                
                # Track tool usage
                tools_used.append(tool_name)
                
                print(f"🔧 Ollama tool call: {tool_name} with args: {tool_args}", flush=True)
                
                # Execute tool
                if tool_name == "search_knowledge_base":
                    try:
                        query = tool_args.get("query", original_question)
                        # Enhance query with original question context if query is too short
                        if len(query.split()) < 3 and len(original_question.split()) > 3:
                            # Combine both for better retrieval
                            enhanced_query = f"{query} {original_question}"
                        else:
                            enhanced_query = query
                        
                        print(f"🔍 Searching with query: '{enhanced_query}'", flush=True)
                        # Use enhanced query for better context
                        docs = fetch_context(enhanced_query, use_mmr=True)
                        
                        # Format results with better context
                        formatted_results = []
                        for doc in docs:
                            formatted_results.append({
                                "content": doc.page_content,
                                "source": doc.metadata.get("source", "unknown"),
                                "metadata": doc.metadata
                            })
                        
                        result = {
                            "query": query,
                            "original_question": original_question,
                            "results": formatted_results,
                            "count": len(docs),
                            "message": f"Found {len(docs)} relevant document chunks. Read them carefully to find the exact answer."
                        }
                        print(f"✅ Tool result: Found {len(docs)} documents", flush=True)
                    except Exception as e:
                        result = {"error": str(e), "message": "Error searching knowledge base"}
                
                # Format tool result for next iteration
                tool_results = json.dumps(result, default=str, ensure_ascii=False)
                
                # Continue to next iteration with tool results
                continue
            else:
                # No tool call - this might be the final answer, but check if we have context
                if tool_results and docs:
                    # We have context, make sure the answer is based on it
                    try:
                        result_data = json.loads(tool_results)
                        if result_data.get("results"):
                            # Verify answer is reasonable, if not, try to extract from context
                            answer_prompt = f"""Based on the following context from uploaded documents, answer this question: "{original_question}"

CONTEXT FROM DOCUMENTS:
{chr(10).join([f"--- Chunk {i+1} ---{chr(10)}{doc.get('content', '')[:400]}" for i, doc in enumerate(result_data.get('results', [])[:3])])}

QUESTION: {original_question}

INSTRUCTIONS:
1. Read the context above carefully
2. Find the EXACT answer to the question
3. If the question asks for a "company name", extract ONLY the company/employer name
4. Answer directly and concisely - just the requested information
5. If the answer is not in the context, say "I couldn't find this information"

Your answer:"""
                            
                            final_msg = [HumanMessage(content=answer_prompt)]
                            final_response = ollama_llm.invoke(final_msg)
                            return final_response.content.strip(), docs
                    except Exception as e:
                        print(f"⚠️  Error in final answer refinement: {e}", flush=True)
                
                # Return the response as-is
                return response_text, docs
                
        except Exception as e:
            print(f"❌ Error in Ollama answer: {e}", flush=True)
            import traceback
            traceback.print_exc()
            return f"Sorry, an error occurred: {str(e)}", docs
    
    # Max iterations reached
    return "Sorry, maximum iterations reached. Please try rephrasing your question.", docs

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document RAG Chatbot</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            height: 100vh;
            margin: 0;
            padding: 20px;
            overflow: hidden;
        }
        
        .main-container {
            display: flex;
            width: 100%;
            height: 100%;
            gap: 20px;
            justify-content: center;
            align-items: center;
        }
        
        .chat-container {
            width: 100%;
            max-width: 600px;
            height: 100%;
            max-height: 90vh;
            background: #f0f2f5;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            display: flex;
            flex-direction: column;
            overflow: hidden;
        }
        
        .chat-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        
        .header-left {
            display: flex;
            align-items: center;
            gap: 15px;
        }
        
        .avatar {
            width: 50px;
            height: 50px;
            border-radius: 50%;
            background: rgba(255,255,255,0.2);
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 24px;
            font-weight: bold;
        }
        
        .header-info h1 {
            font-size: 18px;
            font-weight: 600;
            margin-bottom: 3px;
        }
        
        .header-info p {
            font-size: 13px;
            opacity: 0.9;
        }
        
        .upload-btn {
            background: rgba(255,255,255,0.2);
            border: 1px solid rgba(255,255,255,0.3);
            color: white;
            padding: 8px 16px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 14px;
            transition: all 0.3s;
        }
        
        .upload-btn:hover {
            background: rgba(255,255,255,0.3);
        }
        
        .chat-messages {
            flex: 1;
            overflow-y: auto;
            padding: 20px;
            background: #efeae2;
            background-image: 
                repeating-linear-gradient(0deg, transparent, transparent 2px, rgba(0,0,0,.03) 2px, rgba(0,0,0,.03) 4px);
        }
        
        .chat-messages::-webkit-scrollbar {
            width: 6px;
        }
        
        .chat-messages::-webkit-scrollbar-track {
            background: transparent;
        }
        
        .chat-messages::-webkit-scrollbar-thumb {
            background: rgba(0,0,0,0.2);
            border-radius: 3px;
        }
        
        .message {
            margin-bottom: 15px;
            display: flex;
            animation: slideIn 0.3s ease-out;
        }
        
        @keyframes slideIn {
            from {
                opacity: 0;
                transform: translateY(10px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        .message.user {
            justify-content: flex-end;
        }
        
        .message-bubble {
            max-width: 75%;
            padding: 10px 15px;
            border-radius: 18px;
            word-wrap: break-word;
            box-shadow: 0 1px 2px rgba(0,0,0,0.1);
        }
        
        .message.user .message-bubble {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-bottom-right-radius: 4px;
        }
        
        .message.bot .message-bubble {
            background: white;
            color: #111b21;
            border-bottom-left-radius: 4px;
        }
        
        .message-time {
            font-size: 11px;
            opacity: 0.7;
            margin-top: 5px;
            text-align: right;
        }
        
        .message.bot .message-time {
            text-align: left;
        }
        
        .typing-indicator {
            display: none;
            padding: 10px 15px;
            background: white;
            border-radius: 18px;
            border-bottom-left-radius: 4px;
            max-width: 75px;
        }
        
        .typing-indicator.active {
            display: block;
        }
        
        .typing-dots {
            display: flex;
            gap: 4px;
        }
        
        .typing-dots span {
            width: 8px;
            height: 8px;
            border-radius: 50%;
            background: #999;
            animation: typing 1.4s infinite;
        }
        
        .typing-dots span:nth-child(2) {
            animation-delay: 0.2s;
        }
        
        .typing-dots span:nth-child(3) {
            animation-delay: 0.4s;
        }
        
        @keyframes typing {
            0%, 60%, 100% {
                transform: translateY(0);
                opacity: 0.7;
            }
            30% {
                transform: translateY(-10px);
                opacity: 1;
            }
        }
        
        .chat-input-container {
            background: #f0f2f5;
            padding: 15px;
            border-top: 1px solid rgba(0,0,0,0.1);
        }
        
        .chat-input-wrapper {
            display: flex;
            gap: 10px;
            align-items: flex-end;
            background: white;
            border-radius: 25px;
            padding: 8px 15px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }
        
        .chat-input {
            flex: 1;
            border: none;
            outline: none;
            font-size: 15px;
            font-family: inherit;
            resize: none;
            max-height: 100px;
            padding: 8px 0;
            background: transparent;
        }
        
        .send-button {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            border: none;
            width: 40px;
            height: 40px;
            border-radius: 50%;
            color: white;
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: transform 0.2s, box-shadow 0.2s;
            flex-shrink: 0;
        }
        
        .send-button:hover {
            transform: scale(1.05);
            box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
        }
        
        .send-button:disabled {
            opacity: 0.5;
            cursor: not-allowed;
        }
        
        .welcome-message {
            text-align: center;
            padding: 40px 20px;
            color: #667eea;
        }
        
        .welcome-message h2 {
            font-size: 24px;
            margin-bottom: 10px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        
        .welcome-message p {
            color: #666;
            font-size: 14px;
        }
        
        .upload-modal {
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0,0,0,0.5);
            z-index: 1000;
            align-items: center;
            justify-content: center;
        }
        
        .upload-modal.active {
            display: flex;
        }
        
        .upload-modal-content {
            background: white;
            border-radius: 20px;
            padding: 30px;
            max-width: 500px;
            width: 90%;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
        }
        
        .upload-modal-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
        }
        
        .upload-modal-header h2 {
            color: #667eea;
        }
        
        .close-btn {
            background: none;
            border: none;
            font-size: 24px;
            cursor: pointer;
            color: #999;
        }
        
        .file-upload-area {
            border: 2px dashed #667eea;
            border-radius: 12px;
            padding: 40px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s;
            margin-bottom: 20px;
        }
        
        .file-upload-area:hover {
            background: #f8f9fa;
            border-color: #764ba2;
        }
        
        .file-upload-area.dragover {
            background: #f0f2f5;
            border-color: #764ba2;
        }
        
        .file-input {
            display: none;
        }
        
        .upload-button {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 12px 24px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 16px;
            width: 100%;
            margin-top: 10px;
        }
        
        .upload-button:disabled {
            opacity: 0.5;
            cursor: not-allowed;
        }
        
        .upload-status {
            margin-top: 15px;
            padding: 10px;
            border-radius: 8px;
            display: none;
        }
        
        .upload-status.success {
            background: #d4edda;
            color: #155724;
            display: block;
        }
        
        .upload-status.error {
            background: #f8d7da;
            color: #721c24;
            display: block;
        }
    </style>
</head>
<body>
    <div class="main-container">
        <div class="chat-container">
            <div class="chat-header">
                <div class="header-left">
                    <div class="avatar">🤖</div>
                    <div class="header-info">
                        <h1>Document RAG Assistant</h1>
                        <p>Upload documents and ask questions</p>
                    </div>
                </div>
                <button class="upload-btn" onclick="showUploadModal()">📄 Upload</button>
            </div>
            
            <div class="chat-messages" id="chatMessages">
                <div class="welcome-message">
                    <h2>👋 Welcome!</h2>
                    <p>Upload documents (PDF, Word, TXT) and ask questions about them.</p>
                </div>
            </div>
        
            <div class="typing-indicator" id="typingIndicator">
                <div class="typing-dots">
                    <span></span>
                    <span></span>
                    <span></span>
                </div>
            </div>
            
            <div class="chat-input-container">
                <div class="chat-input-wrapper">
                    <textarea 
                        class="chat-input" 
                        id="messageInput" 
                        placeholder="Ask a question about your documents..."
                        rows="1"
                        onkeydown="handleKeyDown(event)"
                    ></textarea>
                    <button class="send-button" id="sendButton" onclick="sendMessage()">
                        <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                            <line x1="22" y1="2" x2="11" y2="13"></line>
                            <polygon points="22 2 15 22 11 13 2 9 22 2"></polygon>
                        </svg>
                    </button>
                </div>
            </div>
        </div>
    </div>
    
    <div class="upload-modal" id="uploadModal">
        <div class="upload-modal-content">
            <div class="upload-modal-header">
                <h2>Upload Document</h2>
                <button class="close-btn" onclick="hideUploadModal()">×</button>
            </div>
            <div class="file-upload-area" id="fileUploadArea" onclick="document.getElementById('fileInput').click()">
                <p>📄 Click or drag files here</p>
                <p style="font-size: 12px; color: #999; margin-top: 10px;">Supports PDF, Word (.docx), and TXT files</p>
            </div>
            <input type="file" id="fileInput" class="file-input" accept=".pdf,.docx,.doc,.txt" onchange="handleFileSelect(event)">
            <button class="upload-button" id="uploadButton" onclick="uploadFile()" disabled>Upload</button>
            <div class="upload-status" id="uploadStatus"></div>
        </div>
    </div>

    <script>
        let chatHistory = [];
        let selectedFile = null;
        
        function handleKeyDown(event) {
            if (event.key === 'Enter' && !event.shiftKey) {
                event.preventDefault();
                sendMessage();
            }
        }
        
        function autoResize(textarea) {
            textarea.style.height = 'auto';
            textarea.style.height = Math.min(textarea.scrollHeight, 100) + 'px';
        }
        
        document.getElementById('messageInput').addEventListener('input', function() {
            autoResize(this);
        });
        
        function addMessage(content, isUser) {
            const messagesDiv = document.getElementById('chatMessages');
            const welcomeMsg = messagesDiv.querySelector('.welcome-message');
            if (welcomeMsg) {
                welcomeMsg.remove();
            }
            
            const messageDiv = document.createElement('div');
            messageDiv.className = `message ${isUser ? 'user' : 'bot'}`;
            
            const time = new Date().toLocaleTimeString('en-US', { 
                hour: '2-digit', 
                minute: '2-digit' 
            });
            
            messageDiv.innerHTML = `
                <div class="message-bubble">
                    ${content.replace(/\\n/g, '<br>')}
                    <div class="message-time">${time}</div>
                </div>
            `;
            
            messagesDiv.appendChild(messageDiv);
            messagesDiv.scrollTop = messagesDiv.scrollHeight;
        }
        
        function showTyping() {
            const typingIndicator = document.getElementById('typingIndicator');
            const messagesDiv = document.getElementById('chatMessages');
            typingIndicator.classList.add('active');
            messagesDiv.appendChild(typingIndicator);
            messagesDiv.scrollTop = messagesDiv.scrollHeight;
        }
        
        function hideTyping() {
            const typingIndicator = document.getElementById('typingIndicator');
            typingIndicator.classList.remove('active');
        }
        
        async function sendMessage() {
            const input = document.getElementById('messageInput');
            const message = input.value.trim();
            const sendButton = document.getElementById('sendButton');
            
            if (!message) return;
            
            input.disabled = true;
            sendButton.disabled = true;
            
            addMessage(message, true);
            chatHistory.push({ role: 'user', content: message });
            
            input.value = '';
            input.style.height = 'auto';
            
            showTyping();
            
            try {
                const response = await fetch('/api/chat', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({
                        message: message,
                        history: chatHistory
                    })
                });
                
                const data = await response.json();
                
                if (data.error) {
                    addMessage('Sorry, an error occurred: ' + data.error, false);
                } else {
                    addMessage(data.response, false);
                    chatHistory.push({ role: 'assistant', content: data.response });
                }
            } catch (error) {
                addMessage('Connection error: ' + error.message, false);
            } finally {
                hideTyping();
                input.disabled = false;
                sendButton.disabled = false;
                input.focus();
            }
        }
        
        function showUploadModal() {
            document.getElementById('uploadModal').classList.add('active');
        }
        
        function hideUploadModal() {
            document.getElementById('uploadModal').classList.remove('active');
            selectedFile = null;
            document.getElementById('fileInput').value = '';
            document.getElementById('uploadButton').disabled = true;
            document.getElementById('uploadStatus').classList.remove('success', 'error');
            document.getElementById('uploadStatus').style.display = 'none';
        }
        
        function handleFileSelect(event) {
            selectedFile = event.target.files[0];
            if (selectedFile) {
                document.getElementById('uploadButton').disabled = false;
                document.getElementById('fileUploadArea').innerHTML = `
                    <p>📄 ${selectedFile.name}</p>
                    <p style="font-size: 12px; color: #999; margin-top: 10px;">${(selectedFile.size / 1024 / 1024).toFixed(2)} MB</p>
                `;
            }
        }
        
        // Drag and drop
        const fileUploadArea = document.getElementById('fileUploadArea');
        
        fileUploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            fileUploadArea.classList.add('dragover');
        });
        
        fileUploadArea.addEventListener('dragleave', () => {
            fileUploadArea.classList.remove('dragover');
        });
        
        fileUploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            fileUploadArea.classList.remove('dragover');
            const files = e.dataTransfer.files;
            if (files.length > 0) {
                selectedFile = files[0];
                document.getElementById('fileInput').files = files;
                document.getElementById('uploadButton').disabled = false;
                fileUploadArea.innerHTML = `
                    <p>📄 ${selectedFile.name}</p>
                    <p style="font-size: 12px; color: #999; margin-top: 10px;">${(selectedFile.size / 1024 / 1024).toFixed(2)} MB</p>
                `;
            }
        });
        
        async function uploadFile() {
            if (!selectedFile) return;
            
            const uploadButton = document.getElementById('uploadButton');
            const uploadStatus = document.getElementById('uploadStatus');
            
            uploadButton.disabled = true;
            uploadButton.textContent = 'Uploading...';
            uploadStatus.style.display = 'none';
            
            const formData = new FormData();
            formData.append('file', selectedFile);
            
            try {
                const response = await fetch('/api/upload', {
                    method: 'POST',
                    body: formData
                });
                
                const data = await response.json();
                
                if (data.error) {
                    uploadStatus.textContent = 'Error: ' + data.error;
                    uploadStatus.classList.remove('success');
                    uploadStatus.classList.add('error');
                } else {
                    uploadStatus.textContent = `Success! Document "${data.filename}" uploaded and processed. ${data.chunks} chunks added to knowledge base.`;
                    uploadStatus.classList.remove('error');
                    uploadStatus.classList.add('success');
                    selectedFile = null;
                    document.getElementById('fileInput').value = '';
                    
                    setTimeout(() => {
                        hideUploadModal();
                    }, 2000);
                }
            } catch (error) {
                uploadStatus.textContent = 'Error: ' + error.message;
                uploadStatus.classList.remove('success');
                uploadStatus.classList.add('error');
            } finally {
                uploadButton.disabled = false;
                uploadButton.textContent = 'Upload';
            }
        }
        
        window.addEventListener('load', () => {
            document.getElementById('messageInput').focus();
        });
    </script>
</body>
</html>
"""

@app.route('/')
def index():
    """Main chat page."""
    return render_template_string(HTML_TEMPLATE)

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload and process it."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Process the file
        documents = process_uploaded_file(file)
        
        if not documents:
            return jsonify({'error': 'No content extracted from file'}), 400
        
        # Add to vectorstore
        add_documents_to_vectorstore(documents)
        
        # Save file to disk (optional)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.seek(0)  # Reset file pointer
        file.save(file_path)
        
        return jsonify({
            'success': True,
            'filename': file.filename,
            'chunks': len(documents),
            'message': f'Document processed successfully. {len(documents)} chunks added to knowledge base.'
        })
        
    except Exception as e:
        print(f"❌ Error in upload endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat', methods=['POST'])
def chat():
    """Handle chat messages."""
    try:
        data = request.json
        message = data.get('message', '').strip()
        history = data.get('history', [])
        
        if not message:
            return jsonify({'error': 'Message cannot be empty'}), 400
        
        # Get response from RAG system
        response, docs = answer_question_with_ollama(message, history)
        
        # Include source information for debugging
        sources_info = []
        for doc in docs[:3]:  # Top 3 sources
            sources_info.append({
                'source': doc.metadata.get('source', 'unknown'),
                'preview': doc.page_content[:100] + '...' if len(doc.page_content) > 100 else doc.page_content
            })
        
        return jsonify({
            'response': response,
            'sources': len(docs),
            'sources_info': sources_info
        })
    except Exception as e:
        print(f"❌ Error in chat endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("\n" + "=" * 60)
    print("💬 Starting Upload RAG Chatbot Server")
    print("=" * 60)
    print("📍 Open http://127.0.0.1:5001 in your browser")
    print("=" * 60 + "\n")
    
    app.run(host='127.0.0.1', port=5001, debug=False, threaded=True)